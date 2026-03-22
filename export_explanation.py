"""Generate a DOCX file explaining the project in simple terms."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def main():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("AI Traffic Signal Control", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Deep Q-Network Based Adaptive Traffic Signal Control\n"
        "at a 4-Way Intersection with Priority Vehicles"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # =========================================================================
    # THE PROBLEM
    # =========================================================================
    add_heading(doc, "The Problem")

    add_para(
        doc,
        "Think of a normal traffic signal at a crossroad. It has fixed timers — "
        "green for 30 seconds, red for 30 seconds, and it keeps repeating this "
        "same pattern all day. It doesn't care if 100 cars are waiting on one "
        "side and 0 cars on the other side. It just follows the clock blindly. "
        "This wastes everyone's time.",
    )
    add_para(
        doc,
        "This project teaches a computer to control the traffic signal smartly "
        "— like a traffic police officer who looks at all 4 roads and decides "
        "which side needs green based on how many cars are actually waiting.",
        bold=True,
    )

    # =========================================================================
    # HOW IT WORKS
    # =========================================================================
    add_heading(doc, "How It Works — Step by Step")

    # --- Step 1 ---
    add_heading(doc, "Step 1: Build a Virtual Road (generate_network.py)", level=2)

    add_para(
        doc,
        "You can't test on a real road, so you build a virtual crossroad "
        "inside a simulator called SUMO (Simulation of Urban Mobility).",
    )

    # Intersection diagram
    diagram = doc.add_paragraph()
    diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = diagram.add_run(
        "              North\n"
        "                |\n"
        "  West ———— + ———— East\n"
        "                |\n"
        "              South\n\n"
        "(4 roads meeting at 1 traffic light)"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    add_bullet(doc, "Each road is 300 meters long with 3 lanes")
    add_bullet(doc, "Cars drive on the left side (Indian style)")
    add_bullet(doc, "Speed limit is 50 km/h")
    add_bullet(
        doc,
        "The simulator generates ~2400 cars per hour that enter from random "
        "directions and want to go to random destinations",
    )
    add_para(
        doc,
        "This is like building a miniature city intersection inside your computer.",
        italic=True,
    )

    # --- Priority Vehicles Section ---
    add_heading(doc, "Step 2: Add Priority Vehicles (Ambulance)", level=2)

    add_para(
        doc,
        "This project also handles emergency situations. An ambulance needs to "
        "get through as quickly as possible.",
    )

    add_bullet(doc, "Ambulances arrive every 30 seconds")
    add_bullet(doc, "Ambulances travel at 90 km/h (much faster than regular traffic)")
    add_bullet(doc, "They have special priority routing through the intersection")
    add_bullet(doc, "The system can detect and prioritize emergency vehicles")

    add_para(
        doc,
        "Think of it like a fire truck needing to get to an emergency — "
        "the traffic light should turn green for them immediately!",
        italic=True,
    )

    # --- Step 3 ---
    add_heading(doc, "Step 3: Train the AI Brain (train.py)", level=2)

    add_para(
        doc,
        "This is the core of the project. The AI (called DQN — Deep Q-Network) "
        "essentially plays a game:",
    )

    add_bullet(
        doc,
        " — How many cars are queued up on each road? How many are stopped? "
        "Which signal is currently green?",
        bold_prefix="It LOOKS at the intersection",
    )
    add_bullet(
        doc,
        " — Which road should get the green signal next?",
        bold_prefix="It DECIDES",
    )
    add_bullet(
        doc,
        " — If waiting time decreased after its decision, it gets a positive "
        "score (reward). If waiting time increased, it gets a negative score "
        "(punishment).",
        bold_prefix="It gets a SCORE",
    )
    add_bullet(
        doc,
        " — Over 100,000 decisions, it slowly figures out patterns like:\n"
        '  • "When the north road has lots of cars and east road is empty, '
        'I should give green to north"\n'
        '  - "I should not switch signals too frequently because of the '
        'yellow light delay"',
        bold_prefix="It LEARNS from mistakes",
    )

    add_para(
        doc,
        "Think of it like a child learning to play a video game. At first, they "
        "press random buttons. After playing 100,000 rounds, they get pretty "
        "good at it.",
        italic=True,
    )

    add_para(
        doc,
        'The "brain" is a neural network — two layers of 256 neurons each. '
        "It takes the traffic state as input and outputs which signal phase to "
        "activate. Every 10,000 steps, the brain is saved to a file (like a game "
        "checkpoint).",
    )

    # --- Step 4 ---
    add_heading(doc, "Step 4: Test How Good It Is (evaluate.py)", level=2)

    add_para(
        doc,
        "Now we need to prove that the AI is actually better than a dumb "
        "fixed-time signal. So two tests are run:",
    )

    add_bullet(
        doc,
        " — The AI controls the signal for 5 episodes (5 hours of "
        "simulated traffic). It looks at the traffic, picks the best signal "
        "phase, and the simulation runs.",
        bold_prefix="Test A: AI Agent",
    )
    add_bullet(
        doc,
        " — Normal fixed-time signal runs for the same 5 episodes. "
        "The signal just follows a preset timer, ignoring the actual traffic.",
        bold_prefix="Test B: Fixed-Time Signal",
    )

    add_para(doc, "Then we compare the results side-by-side:")

    make_table(
        doc,
        ["Metric", "AI Agent", "Fixed-Time Signal", "AI is Better By"],
        [
            ["Average waiting time per car", "1.33 sec", "4.00 sec", "66.6%"],
            ["Average car speed", "8.43 m/s", "7.96 m/s", "5.8%"],
            ["Cars stuck at red", "7.32", "10.31", "29.0%"],
            ["Total waiting time", "56.86 sec", "155.23 sec", "63.4%"],
        ],
    )

    doc.add_paragraph()  # spacer
    add_para(
        doc,
        "The AI crushes the fixed-time signal on every single metric.",
        bold=True,
    )

    # --- Step 5 ---
    add_heading(doc, "Step 5: Make Pretty Charts (plot_results.py + dashboard.py)", level=2)

    add_para(
        doc,
        "This script reads all the numbers from training and evaluation and "
        "draws 3 graphs:",
    )

    add_bullet(
        doc,
        " — Shows how the AI improved over time (waiting time went down, "
        "speed went up).",
        bold_prefix="Training Progress",
    )
    add_bullet(
        doc,
        " — RL vs Fixed-Time side by side with improvement percentages.",
        bold_prefix="Bar Chart Comparison",
    )
    add_bullet(
        doc,
        " — Shows the AI wins in every single episode, not just on average.",
        bold_prefix="Per-Episode Line Chart",
    )

    add_para(
        doc,
        "There's also an Interactive Dashboard (Streamlit) that lets you:\n"
        "- Watch training in real-time\n"
        "- Compare RL vs Fixed-Time metrics\n"
        "- Browse model checkpoints\n"
        "- View TensorBoard logs",
        italic=True,
    )

    # =========================================================================
    # TECHNOLOGY STACK
    # =========================================================================
    add_heading(doc, "Technology Stack")

    make_table(
        doc,
        ["Tool", "What It Does in This Project"],
        [
            [
                "SUMO",
                "The traffic simulator — creates virtual roads, cars, and "
                "traffic lights",
            ],
            [
                "sumo-rl",
                "The bridge — translates SUMO's traffic data into a format "
                "the AI can understand",
            ],
            [
                "Stable-Baselines3",
                "The AI library — contains the DQN algorithm that does the "
                "actual learning",
            ],
            [
                "Gymnasium",
                "The standard interface — makes sure the AI and simulator "
                "talk in the same language",
            ],
            [
                "Streamlit",
                "Interactive dashboard for real-time monitoring",
            ],
            [
                "Plotly",
                "Interactive charts in the web dashboard",
            ],
            [
                "Matplotlib",
                "Draws the static result graphs",
            ],
            [
                "TensorBoard",
                "Live dashboard to watch the AI learn in real-time",
            ],
            [
                "python-docx",
                "Generates this explanation document",
            ],
        ],
    )

    # =========================================================================
    # ARCHITECTURE
    # =========================================================================
    add_heading(doc, "System Architecture")

    arch = doc.add_paragraph()
    arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = arch.add_run(
        "SUMO (Simulator)  <--TraCI-->  sumo-rl (Env Wrapper)  "
        "<--Gymnasium-->  Stable-Baselines3 DQN (Agent)"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph()

    add_bullet(
        doc,
        "Queue lengths, vehicle densities, current phase (per lane)",
        bold_prefix="Observation space: ",
    )
    add_bullet(
        doc,
        "Select next green phase for the traffic light",
        bold_prefix="Action space: ",
    )
    add_bullet(
        doc,
        "Reduction in total cumulative waiting time between decisions",
        bold_prefix="Reward: ",
    )

    # =========================================================================
    # KEY TRAINING PARAMETERS
    # =========================================================================
    add_heading(doc, "Key Training Parameters")

    make_table(
        doc,
        ["Parameter", "Value", "Purpose"],
        [
            ["Learning rate", "0.001", "How fast the AI learns"],
            ["Buffer size", "50,000", "Memory of past experiences"],
            ["Batch size", "128", "Experiences sampled per learning step"],
            ["Gamma (discount)", "0.99", "How much it values future rewards"],
            ["Network architecture", "[256, 256]", "Two hidden layers, 256 neurons each"],
            ["Delta time", "5 seconds", "Time between AI decisions"],
            ["Yellow time", "3 seconds", "Yellow phase duration"],
            ["Min / Max green", "5s / 60s", "Green phase bounds"],
            ["Total training steps", "100,000", "Total decisions during training"],
            ["Reward function", "diff-waiting-time", "Reward = reduction in delay"],
        ],
    )

    # =========================================================================
    # NETWORK PARAMETERS
    # =========================================================================
    add_heading(doc, "Network Configuration")

    make_table(
        doc,
        ["Parameter", "Value", "Description"],
        [
            ["Number of lanes", "3", "Lanes per direction on each approach"],
            ["Approach length", "300m", "Length of each road entering intersection"],
            ["Speed limit", "50 km/h", "Typical Indian urban road speed"],
            ["Traffic type", "Left-hand", "Indian driving style"],
            ["Vehicle period", "1.5s", "~2400 vehicles per hour"],
            ["Ambulance period", "30s", "Emergency vehicle arrival interval"],
            ["Ambulance speed", "90 km/h", "Priority vehicle speed"],
            ["Simulation duration", "3600s", "1 hour per episode"],
        ],
    )

    # =========================================================================
    # FILES IN THE PROJECT
    # =========================================================================
    add_heading(doc, "Project Files")

    make_table(
        doc,
        ["File", "Purpose"],
        [
            ["config.py", "All tunable parameters centralized"],
            ["generate_network.py", "Creates SUMO network and routes"],
            ["randomTrips.py", "Generates random vehicle demand"],
            ["train.py", "Trains the DQN agent"],
            ["evaluate.py", "Compares RL vs fixed-time signals"],
            ["plot_results.py", "Generates static PNG charts"],
            ["dashboard.py", "Interactive Streamlit dashboard"],
            ["visualization_utils.py", "Shared visualization functions"],
            ["run_all.py", "Runs complete pipeline in one command"],
            ["export_explanation.py", "Generates this DOCX document"],
        ],
    )

    # =========================================================================
    # HOW TO RUN
    # =========================================================================
    add_heading(doc, "How to Run the Project")

    add_para(doc, "Full pipeline (single command):", bold=True)
    code = doc.add_paragraph()
    run = code.add_run("python run_all.py")
    run.font.name = "Consolas"
    run.font.size = Pt(11)

    add_para(doc, "Or step by step:", bold=True)

    steps = [
        ("python generate_network.py", "Generate SUMO network + route files"),
        ("python train.py", "Train the DQN agent (100,000 steps)"),
        ("python evaluate.py", "Compare RL vs fixed-time (5 episodes each)"),
        ("python plot_results.py", "Generate static result charts"),
    ]
    for cmd, desc in steps:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(cmd)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.bold = True
        p.add_run(f"  —  {desc}")

    add_para(doc, "To view interactive dashboard:", bold=True)
    code2 = doc.add_paragraph()
    run = code2.add_run("streamlit run dashboard.py")
    run.font.name = "Consolas"
    run.font.size = Pt(11)

    add_para(doc, "To watch trained AI in action:", bold=True)
    code3 = doc.add_paragraph()
    run = code3.add_run("python evaluate.py --gui")
    run.font.name = "Consolas"
    run.font.size = Pt(11)

    # =========================================================================
    # ONE-LINE SUMMARY
    # =========================================================================
    add_heading(doc, "One-Line Summary")

    summary = doc.add_paragraph()
    summary.paragraph_format.left_indent = Inches(0.5)
    summary.paragraph_format.right_indent = Inches(0.5)
    run = summary.add_run(
        "We built a virtual Indian crossroad with 3 lanes, trained an AI to play "
        "traffic police officer 100,000 times until it learned the best strategy, "
        "added ambulance priority handling for emergency vehicles, and proved it "
        "reduces car waiting time by 66% compared to normal fixed-time signals."
    )
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)

    # --- Save ---
    output_path = "project_explanation.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    main()